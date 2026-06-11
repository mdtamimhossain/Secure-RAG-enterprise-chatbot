from __future__ import annotations

import hashlib
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}
METADATA_SUFFIX = ".metadata.json"

DEFAULT_ROLE_ACCESS = {
    "portfolio": ["visitor"],
    "general": ["employee", "hr", "finance", "manager", "executive", "admin"],
    "hr": ["hr", "manager", "executive", "admin"],
    "finance": ["finance", "executive", "admin"],
    "engineering": ["employee", "manager", "executive", "admin"],
    "executive": ["executive", "admin"],
}


@dataclass(frozen=True)
class LoadedDocument:
    """A document loaded from disk with metadata needed by the RAG pipeline.

    The metadata is intentionally flat because vector databases such as Chroma
    work best with simple scalar metadata values.
    """

    content: str
    metadata: dict[str, str]


def load_documents(data_dir: str | Path) -> list[LoadedDocument]:
    """Load company knowledge documents with access-control metadata.

    Expected layout:

        data/
          hr/
          finance/
          general/

    This loader is only for unstructured company documents that will go into the
    vector database. Future employee/user records should come from the database,
    then be combined with retrieved documents inside the RAG chain.

    Optional sidecar metadata can override inferred values:

        data/hr/leave_policy.md
        data/hr/leave_policy.metadata.json

    Example sidecar:

        {
          "category": "leave_policy",
          "confidentiality": "internal",
          "allowed_roles": ["employee", "hr", "manager"]
        }
    """

    root = Path(data_dir).resolve()
    if not root.exists():
        raise FileNotFoundError(f"Data directory does not exist: {root}")

    documents: list[LoadedDocument] = []
    for path in _iter_document_paths(root):
        text = _load_file(path)
        if not text.strip():
            continue

        metadata = _build_metadata(path, root)
        documents.append(
            LoadedDocument(
                content=text,
                metadata=metadata,
            )
        )

    return documents


def _iter_document_paths(root: Path) -> Iterable[Path]:
    for path in sorted(root.rglob("*")):
        is_metadata_file = path.name.endswith(METADATA_SUFFIX)
        if path.is_file() and not is_metadata_file and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path


def _build_metadata(path: Path, root: Path) -> dict[str, str]:
    relative_path = path.relative_to(root)
    path_parts = [part.lower() for part in relative_path.parts]
    top_level_folder = path_parts[0] if path_parts else "general"
    department = _infer_department(path_parts)

    metadata: dict[str, Any] = {
        "document_id": _document_id(relative_path),
        "source": str(path),
        "relative_path": str(relative_path),
        "filename": path.name,
        "file_type": path.suffix.lower().lstrip("."),
        "scope": "company_document",
        "department": department,
        "category": _infer_category(path, department),
        "confidentiality": _infer_confidentiality(top_level_folder),
        "allowed_roles": DEFAULT_ROLE_ACCESS.get(
            department,
            DEFAULT_ROLE_ACCESS["general"],
        ),
    }

    metadata.update(_load_sidecar_metadata(path))
    metadata["allowed_roles"] = _normalize_roles(metadata.get("allowed_roles"))

    return {key: _metadata_value(value) for key, value in metadata.items()}


def _infer_department(path_parts: list[str]) -> str:
    if not path_parts:
        return "general"

    return path_parts[0]


def _infer_confidentiality(top_level_folder: str) -> str:
    if top_level_folder in {"finance", "executive"}:
        return "restricted"
    return "internal"


def _infer_category(path: Path, department: str = "general") -> str:
    name = path.stem.lower()
    if "knowledge_base" in name:
        return f"{department}_knowledge_base"

    category_keywords = {
        "leave_policy": ["leave", "vacation", "sick"],
        "payroll": ["payroll", "salary", "compensation", "payslip"],
        "benefits": ["benefit", "perk", "insurance"],
        "handbook": ["handbook", "getting_started"],
        "remote_work": ["remote", "hybrid"],
        "helpdesk": ["helpdesk", "support"],
        "communication": ["meeting", "communication"],
        "finance_report": ["10k", "annual", "report", "revenue", "budget"],
        "procurement": ["procurement", "vendor", "purchase"],
        "expense_policy": ["expense", "reimbursement"],
        "revenue_recognition": ["recognition"],
        "onboarding": ["onboarding", "preboarding"],
        "employee_relations": ["relations", "case"],
        "performance": ["performance", "review", "goal"],
        "executive_strategy": ["strategy", "memo", "leadership"],
        "board_metrics": ["board", "metric"],
        "acquisition_risk": ["acquisition", "risk"],
        "device_policy": ["device", "laptop", "equipment"],
        "it_policy": ["system", "security", "access"],
        "career": ["career", "title", "promotion"],
    }

    for category, keywords in category_keywords.items():
        if any(keyword in name for keyword in keywords):
            return category

    return "general"


def _load_sidecar_metadata(path: Path) -> dict[str, Any]:
    sidecar_path = path.with_name(f"{path.stem}{METADATA_SUFFIX}")
    if not sidecar_path.exists():
        return {}

    with sidecar_path.open("r", encoding="utf-8") as file:
        return json.load(file)


def _normalize_roles(value: Any) -> str:
    if isinstance(value, str):
        roles = [role.strip().lower() for role in value.split(",")]
    elif isinstance(value, list):
        roles = [str(role).strip().lower() for role in value]
    else:
        roles = []

    return ",".join(role for role in roles if role)


def _metadata_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return ",".join(str(item) for item in value)
    return str(value)


def _document_id(relative_path: Path) -> str:
    digest = hashlib.sha256(str(relative_path).encode("utf-8")).hexdigest()
    return digest[:16]


def _load_file(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        return _load_docx(path)

    raise ValueError(f"Unsupported document type: {path.suffix}")


def _load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[Page {page_number}]\n{page_text}")
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX loading requires python-docx. Install it with: pip install python-docx"
        ) from exc

    doc = Document(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


if __name__ == "__main__":
    default_data_dir = Path(__file__).resolve().parents[1] / "data"
    loaded = load_documents(default_data_dir)
    print(f"Loaded {len(loaded)} document(s) from {default_data_dir}")
    for document in loaded:
        print(
            f"- {document.metadata['scope']} | "
            f"{document.metadata['department']} | "
            f"{document.metadata['category']} | "
            f"{document.metadata['filename']} | "
            f"roles={document.metadata['allowed_roles']} "
            f"({len(document.content)} characters)"
        )
